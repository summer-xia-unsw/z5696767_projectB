"""Update the Word report to use the redesigned figure suite."""
from __future__ import annotations

from pathlib import Path
import shutil
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = ROOT / "report" / "report.docx"
PNG_DIR = ROOT / "results" / "figures" / "Redesign" / "png"
REPORT_WORKFLOW_COPY = ROOT / "report" / "report_projectb_workflow_data_flow.png"

FIGURE_WIDTH = Inches(6.45)

EXISTING_FIGURE_REPLACEMENTS = [
    "02_fund_growth_of_1.png",
    "02_combined_fund_drawdowns.png",
    "05_report_combined_asset_class_weights.png",
    "02_fund_sharpe_barplot.png",
    "03_sector_sentiment_index.png",
    "03_market_fear_greed_index.png",
    "03_fusion_growth_comparison.png",
    "03_fusion_drawdown_comparison.png",
]


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def set_paragraph_text(doc: Document, prefix: str, text: str) -> None:
    paragraph = find_paragraph(doc, prefix)
    paragraph.clear()
    paragraph.add_run(text)


def replace_picture(paragraph: Paragraph, image_path: Path, width=FIGURE_WIDTH) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=width)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_figure_after(anchor: Paragraph, image_name: str, caption: str, body: str) -> Paragraph:
    image_para = insert_paragraph_after(anchor)
    replace_picture(image_para, PNG_DIR / image_name)
    caption_para = insert_paragraph_after(image_para, caption, "Caption")
    body_para = insert_paragraph_after(caption_para, body, "Normal")
    return body_para


def insert_figure_before(anchor: Paragraph, image_name: str, caption: str, body: str) -> Paragraph:
    image_para = insert_paragraph_before(anchor)
    replace_picture(image_para, PNG_DIR / image_name)
    caption_para = insert_paragraph_before(anchor, caption, "Caption")
    body_para = insert_paragraph_before(anchor, body, "Normal")
    return body_para


def update_appendix_checklist(doc: Document) -> None:
    checklist = None
    for table in doc.tables:
        if table.cell(0, 0).text.strip() == "Course required exhibit":
            checklist = table
            break
    if checklist is None:
        raise ValueError("Appendix checklist table not found")

    updates = {
        "Drawdown figure for at least one fund": "Figure 2 and Figure 10",
        "Sharpe or return-vs-risk barplot": "Figure 4 and Figure 5",
        "Sentiment-index time series for equity sectors": "Figure 6",
        "Fusion before-vs-after table and figure": "Table 4, Figure 9, and Figure 10",
    }
    for row in checklist.rows[1:]:
        key = row.cells[0].text.strip()
        if key in updates:
            row.cells[1].text = updates[key]

    additions = [
        ("Sector sentiment ranking and coverage", "Table 3 and Figure 7"),
        ("App journey and precomputed-data map", "Table 5 and Figure 11"),
        ("Workflow and reproducibility diagram", "Appendix Figure A1"),
    ]
    existing = {row.cells[0].text.strip() for row in checklist.rows}
    for left, right in additions:
        if left not in existing:
            row = checklist.add_row()
            row.cells[0].text = left
            row.cells[1].text = right


def clean_unused_docx_media(path: Path) -> int:
    """Remove old image files that are no longer referenced by document.xml."""
    tmp = path.with_suffix(".cleaned.docx")
    with ZipFile(path, "r") as zin:
        document_xml = zin.read("word/document.xml")
        rels_xml = zin.read("word/_rels/document.xml.rels")
        document_root = etree.fromstring(document_xml)
        rels_root = etree.fromstring(rels_xml)
        ns = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        used_rids = set(document_root.xpath(".//a:blip/@r:embed", namespaces=ns))
        remove_targets: set[str] = set()
        for rel in list(rels_root):
            rel_id = rel.get("Id")
            rel_type = rel.get("Type") or ""
            target = rel.get("Target") or ""
            if rel_type.endswith("/image") and rel_id not in used_rids:
                remove_targets.add("word/" + target if not target.startswith("/") else target.lstrip("/"))
                rels_root.remove(rel)
        new_rels_xml = etree.tostring(
            rels_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )
        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in remove_targets:
                    continue
                if item.filename == "word/_rels/document.xml.rels":
                    zout.writestr(item, new_rels_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))
    shutil.move(str(tmp), str(path))
    return len(remove_targets)


def main() -> None:
    doc = Document(REPORT_PATH)

    image_paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//w:drawing")]
    if len(image_paragraphs) != 8:
        raise ValueError(f"Expected 8 existing report images, found {len(image_paragraphs)}")
    for paragraph, image_name in zip(image_paragraphs, EXISTING_FIGURE_REPLACEMENTS):
        replace_picture(paragraph, PNG_DIR / image_name)

    set_paragraph_text(
        doc,
        "NovaAlloc builds a menu",
        "NovaAlloc builds a menu of 12 investable out-of-sample funds across equity-only, crypto-only, and combined equity-plus-crypto universes. The evidence is now read through the redesigned figure sequence: Table 1 ranks the funds, Figures 1-5 connect growth, drawdown, asset allocation, Sharpe ranking, and risk-return positioning, and the later figures test whether news sentiment improves the base portfolio.",
    )
    set_paragraph_text(
        doc,
        "Table 1 establishes the first ranking",
        "Table 1 establishes the first ranking. NovaAlloc beats the required minimum because it does not stop at one combined fund with two methods; it creates 12 funds across three universes and four methods. The best risk-adjusted result is Equity Equal Weight, while Combined Risk Parity is the most defensible balanced product. Figures 1-5 then explain why this ranking is not just a return ranking.",
    )
    set_paragraph_text(
        doc,
        "Figure 1. Growth of USD 1",
        "Figure 1. Growth of USD 1 for NovaAlloc funds, out-of-sample period 2021-2023. The redesigned figure uses Equity, Combined, and Crypto small multiples plus a return snapshot sorted by ending value of USD 1.",
    )
    set_paragraph_text(
        doc,
        "Figure 1 turns Table 1",
        "Figure 1 turns Table 1 into a path that an investor can understand. Crypto Risk Parity and Crypto Equal Weight finish highest, but their path is much more volatile. The small-multiple layout makes the lower-volatility equity and combined funds visible instead of letting crypto dominate the scale.",
    )
    set_paragraph_text(
        doc,
        "Figure 2. Drawdowns",
        "Figure 2. Drawdowns for combined equity-plus-crypto funds, out-of-sample period 2021-2023. The redesigned figure highlights maximum drawdown and reports a right-side risk snapshot for each combined method.",
    )
    set_paragraph_text(
        doc,
        "Figure 2 shows that",
        "Figure 2 shows that the combined universe is not automatically safe. Combined Equal Weight has strong final growth, but its drawdown reaches about 29 percent. Minimum Variance limits the drawdown to about 18 percent, while Maximum Sharpe suffers the deepest loss because expected-return estimation creates a fragile allocation.",
    )
    set_paragraph_text(
        doc,
        "Figure 3. Combined funds' asset-class weights",
        "Figure 3. Combined funds' asset-class weights over time by method, 2021-2023. The redesigned figure compares Equity and Crypto buckets across all four combined-fund methods and reports a crypto-allocation snapshot.",
    )
    set_paragraph_text(
        doc,
        "Figure 3 links the risk result",
        "Figure 3 links the risk result back to portfolio construction. Equal Weight mechanically keeps crypto higher. Minimum Variance almost removes crypto because crypto volatility dominates the covariance matrix. Risk Parity keeps crypto as a small satellite exposure, which explains why it is a credible balanced default rather than a pure return-chasing strategy.",
    )
    set_paragraph_text(
        doc,
        "Figure 4. Sharpe ratios",
        "Figure 4. Sharpe ratios across funds and methods, out-of-sample period 2021-2023. The redesigned ranking highlights the best and worst funds and reports the top-five Sharpe results in a side panel.",
    )
    set_paragraph_text(
        doc,
        "Figure 4 closes the fund comparison",
        "Figures 4 and 5 close the fund comparison. Figure 4 shows that Equal Weight and Risk Parity are stronger out-of-sample baselines than the estimated Maximum Sharpe optimiser in this sample. This result is economically plausible because Maximum Sharpe relies heavily on noisy expected returns, while Equal Weight and Risk Parity depend more on robust diversification rules.",
    )

    anchor = find_paragraph(doc, "Figures 4 and 5 close")
    insert_figure_after(
        anchor,
        "02_fund_risk_return_scatter.png",
        "Figure 5. Risk-return map for all NovaAlloc funds, 2021-2023. Annualised return is plotted against annualised volatility, with dotted zero-rate Sharpe guide lines.",
        "Figure 5 explains why the Sharpe result is not simply a ranking of returns. Crypto Risk Parity has the highest raw annualised return, but it sits far to the right of the chart and carries very deep drawdown risk. Equity Equal Weight and Combined Risk Parity are more useful product candidates because they convert risk into return more consistently.",
    )

    set_paragraph_text(
        doc,
        "Table 3 checks the sentiment data",
        "Table 3 checks the sentiment data before the time-series figures are interpreted. Average sector scores are above 50 in every high-coverage sector reported here, which means the news stream is mildly positive on average. Figures 6 and 7 separate two questions: whether sentiment moves through time, and which sectors have the highest average tone.",
    )
    set_paragraph_text(
        doc,
        "Figure 5. Sector fear/greed",
        "Figure 6. Sector fear/greed sentiment index over time, 2020-2023. The redesigned figure uses sector small multiples to show 21-day rolling 0-100 scores without a crowded ten-line legend.",
    )
    set_paragraph_text(
        doc,
        "Figure 5 shows why",
        "Figure 6 shows why the sector index belongs in the app even before it is used for trading. Sentiment changes over time and sectors do not move identically. The index should not be sold as a clean return predictor because headlines are short and noisy, but it is useful context for interpreting investor mood and sector-specific news pressure.",
    )
    anchor = find_paragraph(doc, "Figure 6 shows why")
    insert_figure_after(
        anchor,
        "03_sector_sentiment_ranking.png",
        "Figure 7. Average sector fear/greed ranking, 2020-2023. The redesigned ranking compares full-sample sector sentiment averages and headline coverage.",
        "Figure 7 makes the cross-sectional sentiment result easier to read. Utilities, Real Estate and Technology have the highest average fear/greed scores, while Financials is the lowest among the reported sectors. This ranking is descriptive evidence, not a trading signal by itself, because the portfolio overlay must use lagged sector signals.",
    )
    set_paragraph_text(
        doc,
        "Figure 6. Market-level news",
        "Figure 8. Market-level news fear/greed index, 2020-2023. The redesigned figure shows the 21-day rolling market index, the neutral 50 line, and a right-side market snapshot.",
    )
    set_paragraph_text(
        doc,
        "Figure 6 gives the user",
        "Figure 8 gives the user a simpler market fear/greed reading across 1,006 trading days from 2020-01-02 to 2023-12-29. The market index is useful because the app user can compare fund performance with the tone of recent equity news. At the same time, the market index is deliberately not used as a same-day trading signal, because that would create look-ahead risk.",
    )
    set_paragraph_text(
        doc,
        "Table 4 gives the main result",
        "Table 4 gives the main result: the sentiment tilt reduces annualised return from 7.0% to 6.2% and reduces Sharpe from 0.550 to 0.482. Maximum drawdown improves slightly from -18.3% to -17.9%, but the smaller drawdown is not enough to offset the lower return. The diagnostic check reports zero lag violations, a mean absolute weight change of about 0.17 percentage points, and a maximum absolute weight change of about 4.9 percentage points. This means the test is not look-ahead biased and the tilt is economically modest. Figures 9 and 10 then show whether these table differences are visible in the return path.",
    )

    set_paragraph_text(
        doc,
        "Figure 7. Fusion growth",
        "Figure 9. Fusion growth comparison: base equity minimum-variance fund versus sentiment tilt, 2021-2023. The redesigned figure highlights that the sentiment tilt finishes below the base fund.",
    )
    set_paragraph_text(
        doc,
        "Figure 7 shows that",
        "Figure 9 shows that the sentiment tilt does not create a better growth path in this sample. The correct interpretation is not that sentiment has no value. The result says that this simple linear momentum tilt is not strong enough to market as alpha after it is made lagged and investable.",
    )
    set_paragraph_text(
        doc,
        "Figure 8. Fusion drawdown",
        "Figure 10. Fusion drawdown comparison: base equity minimum-variance fund versus sentiment tilt, 2021-2023. The redesigned figure shows the small drawdown improvement next to the lower-growth result.",
    )
    set_paragraph_text(
        doc,
        "Figure 8 shows a small",
        "Figure 10 shows a small drawdown benefit from the sentiment tilt. A risk-focused investor may value this direction, but the improvement is too small to justify the overlay without transaction costs, more validation, and tests over different market regimes. This is why the report treats sentiment as useful analytics first and as a trading overlay second.",
    )

    set_paragraph_text(
        doc,
        "The app journey follows",
        "The app journey follows the same order as the redesigned report. Overview starts with the product and the best fund evidence. Funds lets the user filter by universe and method, compare metrics, read a fund fact sheet, and inspect current holdings. Allocation explains portfolio weights. Sentiment shows the sector and market fear/greed views. Fusion reports the before-vs-after result. Data Health shows which precomputed files were loaded.",
    )
    set_paragraph_text(
        doc,
        "The design choice to move navigation",
        "The design choice to move navigation into the sidebar is practical. Top tabs were harder to read on narrower screens, while the sidebar keeps the full investor journey visible. The app also includes a data health page and manifest files so that a marker can verify the required artifacts without searching through raw CSV files.",
    )
    anchor = find_paragraph(doc, "The design choice to move navigation")
    insert_figure_after(
        anchor,
        "04_app_page_data_map.png",
        "Figure 11. Streamlit app page-data map. The redesigned diagram shows how precomputed results feed the six app pages without rerunning backtests or VADER at runtime.",
        "Figure 11 turns the app requirement into a clear product journey. It also documents the deployment logic: the app consumes precomputed results, which keeps the interface light and reduces the risk that a public deployment fails because of a long backtest or sentiment rebuild.",
    )

    set_paragraph_text(
        doc,
        "The main innovation is the way",
        "The main innovation is the way the project links the stages together, not a single complicated model. The required minimum is a combined equity-plus-crypto fund with two methods. NovaAlloc builds 12 funds across three universes and four methods, creates a sector sentiment layer, tests a lagged fusion overlay, and turns the result into a Streamlit app. Appendix Figure A1 summarises this end-to-end pipeline.",
    )
    set_paragraph_text(
        doc,
        "The reflection follows the evidence chain",
        "The reflection follows the evidence chain. The fund results are stronger than the first sentiment overlay: simple and risk-focused methods are more reliable than Maximum Sharpe optimisation, and Combined Risk Parity is the most credible default multi-asset product. The redesigned Figures 1-5 make this conclusion visible from both the performance path and the risk-return map.",
    )
    set_paragraph_text(
        doc,
        "Recommendation 1:",
        "Recommendation 1: make Combined Risk Parity the default balanced product, with Combined Equal Weight as the higher-growth option and crypto-only funds labelled as high-risk satellite products. This follows Table 1 and Figures 1-5: combined funds offer a better investor experience than pure crypto even when crypto has higher raw upside.",
    )
    set_paragraph_text(
        doc,
        "AI workflow evidence is stored",
        "AI workflow evidence is stored in AGENTS.md and the numbered ai/ prompt logs. The report and app are based on the precomputed results files listed in Table 5, while Appendix Figure A1 documents how the data foundation, portfolios, sentiment, fusion, app, report and AI workflow checks connect.",
    )

    update_appendix_checklist(doc)

    appendix_b = find_paragraph(doc, "Appendix B. Reproducibility notes")
    insert_figure_before(
        appendix_b,
        "05_report_projectb_workflow_data_flow.png",
        "Appendix Figure A1. NovaAlloc ProjectB workflow and data-flow map. The diagram summarises how raw price/news inputs flow into returns, funds, sentiment, fusion, app, report, and AI workflow evidence.",
        "Appendix Figure A1 is included as reproducibility support rather than as a new modelling result. It helps explain how the report, Streamlit app, and AI logs all use the same canonical results files.",
    )

    doc.save(REPORT_PATH)
    removed_media = clean_unused_docx_media(REPORT_PATH)
    shutil.copy2(PNG_DIR / "05_report_projectb_workflow_data_flow.png", REPORT_WORKFLOW_COPY)
    print(f"Updated {REPORT_PATH}")
    print(f"Updated {REPORT_WORKFLOW_COPY}")
    print(f"Removed {removed_media} unused media files")


if __name__ == "__main__":
    main()
